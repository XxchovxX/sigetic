from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
MANUALS_DIR = ROOT / "docs" / "manuales"
OUTPUT_DIR = MANUALS_DIR / "generados"
DOCX_DIR = OUTPUT_DIR / "word"
PDF_DIR = OUTPUT_DIR / "pdf"
LOGO_PATH = ROOT / "sigetic-web" / "public" / "identity" / "logo-alcaldia.png"

GREEN = "007A33"
DARK_GREEN = "005C2A"
NAVY = "071A3A"
MUTED = "4F6685"
LIGHT_GREEN = "EAF8EF"
BORDER = "222222"


def hex_color(value: str):
    return colors.HexColor("#" + value.lstrip("#"))


@dataclass(frozen=True)
class ManualSpec:
    source: str
    code: str
    title: str


MANUALS: list[ManualSpec] = [
    ManualSpec("00-manual-general-sigetic.md", "TIC-MAN-GEN-01", "MANUAL GENERAL SIGETIC"),
    ManualSpec("01-administrador-general.md", "TIC-MAN-ADM-01", "MANUAL ADMINISTRADOR GENERAL"),
    ManualSpec("02-administrador-tic.md", "TIC-MAN-ATIC-01", "MANUAL ADMINISTRADOR TIC"),
    ManualSpec("03-auxiliar-sistemas.md", "TIC-MAN-ASIS-01", "MANUAL AUXILIAR DE SISTEMAS"),
    ManualSpec("04-secretario-administrativo-financiero.md", "TIC-MAN-SAF-01", "MANUAL SECRETARIO ADMINISTRATIVO FINANCIERO"),
    ManualSpec("05-auxiliar-administrativo-saf.md", "TIC-MAN-AUXSAF-01", "MANUAL AUXILIAR ADMINISTRATIVO SAF"),
    ManualSpec("06-secretario-despacho.md", "TIC-MAN-SDES-01", "MANUAL SECRETARIO DE DESPACHO"),
    ManualSpec("07-funcionario.md", "TIC-MAN-FUNC-01", "MANUAL FUNCIONARIO"),
    ManualSpec("08-consulta-control-interno.md", "TIC-MAN-CI-01", "MANUAL CONSULTA / CONTROL INTERNO"),
    ManualSpec("09-pendientes-integraciones.md", "TIC-MAN-INT-01", "INTEGRACIONES Y CONTROLES PENDIENTES"),
]


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = text.replace("`", "")
    return text.strip()


def parse_markdown(path: Path) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line:
            blocks.append(("blank", ""))
            continue
        if line.startswith("# "):
            blocks.append(("h1", clean_inline(line[2:])))
        elif line.startswith("## "):
            blocks.append(("h2", clean_inline(line[3:])))
        elif line.startswith("### "):
            blocks.append(("h3", clean_inline(line[4:])))
        elif line.startswith("- "):
            blocks.append(("bullet", clean_inline(line[2:])))
        elif re.match(r"^\d+\.\s+", line):
            blocks.append(("number", clean_inline(re.sub(r"^\d+\.\s+", "", line))))
        elif line.startswith("  - "):
            blocks.append(("bullet", clean_inline(line[4:])))
        else:
            blocks.append(("p", clean_inline(line)))
    return compact_blocks(blocks)


def compact_blocks(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    compacted: list[tuple[str, str]] = []
    previous_blank = False
    for kind, text in blocks:
        if kind == "blank":
            if not previous_blank:
                compacted.append((kind, text))
            previous_blank = True
        else:
            compacted.append((kind, text))
            previous_blank = False
    return compacted


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9, color: str = NAVY, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_width(cell, width: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def set_paragraph_font(paragraph, name: str = "Arial", size: int = 11, color: str = NAVY, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:ascii"), name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_docx_header(section, spec: ManualSpec) -> None:
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=4, cols=3, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    widths = [1700, 4800, 2860]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo_cell = table.cell(0, 0).merge(table.cell(3, 0))
    title_cell = table.cell(0, 1).merge(table.cell(3, 1))
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_logo.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(0.95))

    set_cell_text(title_cell, spec.title, bold=True, size=11, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    metadata = [
        f"CÓDIGO: {spec.code}",
        "VERSIÓN: 01",
        "FECHA: JUNIO DE 2026",
        "PÁGINA: ",
    ]
    for idx, value in enumerate(metadata):
        cell = table.cell(idx, 2)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        r.font.size = Pt(8)
        if idx == 3:
            add_field(p, "PAGE")
            p.add_run(" DE ")
            add_field(p, "NUMPAGES")
            set_paragraph_font(p, "Arial", 8, NAVY, True)


def add_docx_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("DESARROLLADO POR ING. CRISTIAN HUMBERTO OVALLE VARÓN")
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ALCALDÍA DE SAN CARLOS DE GUAROA")
    r2.bold = True
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor.from_string(DARK_GREEN)


def configure_docx(doc: Document, spec: ManualSpec) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.55)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    add_docx_header(section, spec)
    add_docx_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 14, 8),
        ("Heading 2", 13, GREEN, 12, 5),
        ("Heading 3", 11, MUTED, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_control_table(doc: Document, spec: ManualSpec) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "A7B6C8", "6")
    rows = [
        ("Documento", spec.title.title()),
        ("Código", spec.code),
        ("Versión", "01"),
        ("Fecha", "Junio de 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], LIGHT_GREEN)
        set_cell_width(row.cells[0], 1800)
        set_cell_width(row.cells[1], 7560)
        set_cell_text(row.cells[0], label, bold=True, size=9, color=DARK_GREEN)
        set_cell_text(row.cells[1], value, size=9, color=NAVY)
    doc.add_paragraph()


def add_docx_content(doc: Document, blocks: Iterable[tuple[str, str]], spec: ManualSpec) -> None:
    seen_title = False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(spec.title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_control_table(doc, spec)

    for kind, text in blocks:
        if kind == "h1" and not seen_title:
            seen_title = True
            continue
        if kind == "blank":
            continue
        if kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif kind == "h3":
            doc.add_paragraph(text, style="Heading 3")
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(text, style="List Number")
        elif kind == "p":
            doc.add_paragraph(text)


def build_docx(spec: ManualSpec) -> Path:
    doc = Document()
    configure_docx(doc, spec)
    blocks = parse_markdown(MANUALS_DIR / spec.source)
    add_docx_content(doc, blocks, spec)
    out = DOCX_DIR / spec.source.replace(".md", ".docx")
    doc.save(out)
    return out


class HeaderFooterCanvas(canvas.Canvas):
    def __init__(self, *args, spec: ManualSpec, **kwargs):
        super().__init__(*args, **kwargs)
        self.spec = spec
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        page_count = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_header_footer(page_count)
            super().showPage()
        super().save()

    def draw_header_footer(self, page_count: int):
        page_no = self._pageNumber
        width, height = letter
        x = 0.72 * inch
        y = height - 0.78 * inch
        table_w = 6.55 * inch
        left_w = 1.08 * inch
        mid_w = 3.4 * inch
        right_w = table_w - left_w - mid_w
        header_h = 0.82 * inch
        row_h = header_h / 4

        self.setStrokeColor(colors.black)
        self.setLineWidth(0.7)
        self.rect(x, y - header_h, table_w, header_h)
        self.line(x + left_w, y, x + left_w, y - header_h)
        self.line(x + left_w + mid_w, y, x + left_w + mid_w, y - header_h)
        for idx in range(1, 4):
            yy = y - idx * row_h
            self.line(x + left_w + mid_w, yy, x + table_w, yy)

        try:
            self.drawImage(
                ImageReader(str(LOGO_PATH)),
                x + 0.17 * inch,
                y - 0.76 * inch,
                width=0.74 * inch,
                height=0.64 * inch,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception:
            pass

        self.setFont("Helvetica-Bold", 9.5)
        self.setFillColor(hex_color(NAVY))
        self.drawCentredString(x + left_w + mid_w / 2, y - 0.47 * inch, self.spec.title)

        meta = [
            f"CÓDIGO: {self.spec.code}",
            "VERSIÓN: 01",
            "FECHA: JUNIO DE 2026",
            f"PÁGINA: {page_no} DE {page_count}",
        ]
        self.setFont("Helvetica-Bold", 7.6)
        for idx, text in enumerate(meta):
            self.drawString(x + left_w + mid_w + 0.08 * inch, y - (idx + 0.68) * row_h, text)

        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(hex_color(MUTED))
        self.drawCentredString(width / 2, 0.47 * inch, "DESARROLLADO POR ING. CRISTIAN HUMBERTO OVALLE VARÓN")
        self.setFillColor(hex_color(DARK_GREEN))
        self.drawCentredString(width / 2, 0.31 * inch, "ALCALDÍA DE SAN CARLOS DE GUAROA")


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ManualTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        textColor=hex_color(NAVY),
        alignment=TA_CENTER,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="H1Manual",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        textColor=hex_color(NAVY),
        spaceBefore=12,
        spaceAfter=7,
    ))
    styles.add(ParagraphStyle(
        name="H2Manual",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        leading=14,
        textColor=hex_color(GREEN),
        spaceBefore=10,
        spaceAfter=5,
    ))
    styles.add(ParagraphStyle(
        name="H3Manual",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=hex_color(MUTED),
        spaceBefore=8,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="BodyManual",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.6,
        leading=12.5,
        textColor=hex_color(NAVY),
        spaceAfter=5,
        alignment=TA_LEFT,
    ))
    styles.add(ParagraphStyle(
        name="BulletManual",
        parent=styles["BodyManual"],
        leftIndent=18,
        firstLineIndent=-9,
        bulletIndent=8,
    ))
    return styles


def build_pdf(spec: ManualSpec) -> Path:
    out = PDF_DIR / spec.source.replace(".md", ".pdf")
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(out),
        pagesize=letter,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=1.78 * inch,
        bottomMargin=0.78 * inch,
    )

    story = [Paragraph(spec.title, styles["ManualTitle"])]
    control = Table(
        [
            ["Documento", spec.title.title()],
            ["Código", spec.code],
            ["Versión", "01"],
            ["Fecha", "Junio de 2026"],
        ],
        colWidths=[1.25 * inch, 5.15 * inch],
    )
    control.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, hex_color("A7B6C8")),
        ("BACKGROUND", (0, 0), (0, -1), hex_color(LIGHT_GREEN)),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.6),
        ("TEXTCOLOR", (0, 0), (0, -1), hex_color(DARK_GREEN)),
        ("TEXTCOLOR", (1, 0), (1, -1), hex_color(NAVY)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story += [control, Spacer(1, 10)]

    blocks = parse_markdown(MANUALS_DIR / spec.source)
    seen_title = False
    pending_list: list[str] = []
    pending_kind: str | None = None

    def flush_list():
        nonlocal pending_list, pending_kind
        if not pending_list:
            return
        if pending_kind == "bullet":
            for item in pending_list:
                story.append(Paragraph(f"&bull;&nbsp;{item}", styles["BulletManual"]))
        else:
            for idx, item in enumerate(pending_list, start=1):
                story.append(Paragraph(f"{idx}.&nbsp;{item}", styles["BulletManual"]))
        story.append(Spacer(1, 3))
        pending_list = []
        pending_kind = None

    for kind, text in blocks:
        if kind == "h1" and not seen_title:
            seen_title = True
            continue
        if kind in {"bullet", "number"}:
            if pending_kind and pending_kind != kind:
                flush_list()
            pending_kind = kind
            pending_list.append(text)
            continue
        flush_list()
        if kind == "blank":
            story.append(Spacer(1, 2))
        elif kind == "h1":
            story.append(Paragraph(text, styles["H1Manual"]))
        elif kind == "h2":
            story.append(Paragraph(text, styles["H2Manual"]))
        elif kind == "h3":
            story.append(Paragraph(text, styles["H3Manual"]))
        elif kind == "p":
            story.append(Paragraph(text, styles["BodyManual"]))
    flush_list()

    doc.build(story, canvasmaker=lambda *args, **kwargs: HeaderFooterCanvas(*args, spec=spec, **kwargs))
    return out


def main() -> None:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    if not LOGO_PATH.exists():
        raise FileNotFoundError(f"No existe el logo: {LOGO_PATH}")
    Image.open(LOGO_PATH).verify()

    created: list[Path] = []
    for spec in MANUALS:
        created.append(build_docx(spec))
        created.append(build_pdf(spec))

    print("Documentos generados:")
    for path in created:
        print(path.relative_to(ROOT))


if __name__ == "__main__":
    main()
